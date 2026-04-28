"""
findings.py
-----------
Findings generation and Word export for the Protiviti Operational Audit Assistant.

Generates 5C-format findings from completed work program rows.
Exports findings to Word (.docx).
"""

import json
import logging
import re
import uuid
from datetime import datetime, timezone
from pathlib import Path

from core.llm_client import get_llm_client, get_llm_model, build_completion_kwargs, build_response_format_kwargs
from core.token_tracker import record_usage

logger = logging.getLogger(__name__)


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def generate_findings(audit_id: str, db_conn) -> list[str]:
    """
    Generate findings from completed work program rows.
    Returns list of new finding IDs.
    """
    rows = db_conn.execute(
        """SELECT * FROM work_program_rows
           WHERE audit_id=? AND status IN ('completed', 'open_questions')
           AND conclusion IS NOT NULL AND conclusion != ''""",
        (audit_id,)
    ).fetchall()

    if not rows:
        return []

    client = get_llm_client()
    model = get_llm_model()
    new_ids = []

    for row in rows:
        prompt = f"""You are a senior internal auditor writing an audit finding.

TEST: {row['description']}
CONCLUSION: {row['conclusion']}
EVIDENCE: {row['evidence_references'] or '[]'}
OPEN QUESTIONS: {row['open_questions'] or '[]'}

Write a formal audit finding in 5C format. Return JSON with:
- "title": concise finding title (one sentence)
- "condition": what was observed
- "criteria": the standard or expectation that was not met
- "cause": root cause of the issue
- "consequence": impact or risk if not addressed
- "corrective_action": recommended remediation
- "sub_issues": array of strings for any sub-issues (empty array if none)

Return ONLY valid JSON. If no finding is warranted (test passed cleanly), return null."""

        resp = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            **build_completion_kwargs(max_tokens=1200, temperature=0.2),
            **build_response_format_kwargs(),
        )

        record_usage(audit_id, "findings", resp.usage, db_conn)

        raw = resp.choices[0].message.content or "null"
        try:
            finding = json.loads(raw)
        except json.JSONDecodeError:
            match = re.search(r"\{.*\}", raw, re.DOTALL)
            finding = json.loads(match.group(0)) if match else None

        if not finding:
            continue

        finding_id = str(uuid.uuid4())
        db_conn.execute(
            """INSERT INTO findings
               (id, audit_id, title, condition, criteria, cause, consequence,
                corrective_action, sub_issues, source_row_ids, created_at, updated_at)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (finding_id, audit_id,
             finding.get("title", "Untitled Finding"),
             finding.get("condition", ""),
             finding.get("criteria", ""),
             finding.get("cause", ""),
             finding.get("consequence", ""),
             finding.get("corrective_action", ""),
             json.dumps(finding.get("sub_issues", [])),
             json.dumps([row["id"]]),
             _now(), _now())
        )
        new_ids.append(finding_id)

    db_conn.commit()
    return new_ids


def export_findings_docx(audit_id: str, db_conn, output_path: str) -> str:
    """
    Export all findings for an audit to a Word document.
    Returns the output file path.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    audit = db_conn.execute("SELECT name, client FROM audits WHERE id=?", (audit_id,)).fetchone()
    findings = db_conn.execute(
        "SELECT * FROM findings WHERE audit_id=? ORDER BY created_at",
        (audit_id,)
    ).fetchall()

    doc = Document()

    # Title
    title = doc.add_heading("Audit Findings Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"{audit['name']}")
    run.bold = True
    run.font.size = Pt(14)
    if audit["client"]:
        doc.add_paragraph(f"Client: {audit['client']}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated: {_now()[:10]}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    if not findings:
        doc.add_paragraph("No findings have been generated for this audit.")
    else:
        for i, f in enumerate(findings, 1):
            doc.add_heading(f"Finding {i}: {f['title']}", level=1)

            sections = [
                ("Condition", f["condition"]),
                ("Criteria", f["criteria"]),
                ("Cause", f["cause"]),
                ("Consequence", f["consequence"]),
                ("Corrective Action", f["corrective_action"]),
            ]

            for label, content in sections:
                if content:
                    p = doc.add_paragraph()
                    run = p.add_run(f"{label}: ")
                    run.bold = True
                    p.add_run(content)

            sub_issues = json.loads(f["sub_issues"] or "[]")
            if sub_issues:
                doc.add_paragraph("Sub-Issues:", style="List Bullet")
                for sub in sub_issues:
                    doc.add_paragraph(sub, style="List Bullet")

            doc.add_paragraph("")  # spacing

    doc.save(output_path)
    return output_path
